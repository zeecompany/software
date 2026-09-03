"""Practical PDF toolbox for AURCO's built-in viewer.

This module focuses on fast, offline document operations using libraries already
shipped with the desktop app. The goal is a professional daily-work PDF tool,
not a full Adobe clone.
"""
from __future__ import annotations

import datetime as _dt
import hashlib
import html
import io
import json
import shutil
from pathlib import Path
from typing import Any, Iterable

import pypdfium2 as pdfium
from openpyxl import Workbook
from pypdf import PdfReader, PdfWriter
from pypdf.constants import UserAccessPermissions as UAP
from reportlab.lib import colors
from reportlab.lib.pagesizes import portrait
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph

from . import config

try:
    from docx import Document as DocxDocument
except Exception:  # noqa: BLE001
    DocxDocument = None

PDF_STATE_FILE = config.appdata_dir() / "pdf_viewer_state.json"
PDF_WORK_DIR = config.appdata_dir() / "pdf_viewer_work"
PDF_EXPORT_DIR = config.appdata_dir() / "pdf_viewer_exports"

_STYLES = getSampleStyleSheet()


def _now_tag() -> str:
    return _dt.datetime.now().strftime("%Y%m%d_%H%M%S")


def _safe(name: str, fallback: str = "file") -> str:
    out = "".join(c for c in str(name or "") if c.isalnum() or c in " -_().[]")
    out = " ".join(out.split()).strip(" .")
    return out or fallback


def _read_state() -> dict[str, Any]:
    if PDF_STATE_FILE.exists():
        try:
            return json.loads(PDF_STATE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"recent": [], "sessions": {}}


def _write_state(data: dict[str, Any]) -> None:
    PDF_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
    PDF_STATE_FILE.write_text(json.dumps(data, indent=2), encoding="utf-8")


def _file_key(path: str | Path) -> str:
    return hashlib.sha256(str(Path(path).resolve()).encode("utf-8")).hexdigest()[:24]


def recent_files(limit: int = 14) -> list[str]:
    st = _read_state()
    out = []
    for p in st.get("recent", []):
        if Path(p).exists() and p not in out:
            out.append(p)
        if len(out) >= limit:
            break
    return out


def remember_recent(path: str | Path, limit: int = 14) -> None:
    p = str(Path(path).resolve())
    st = _read_state()
    items = [x for x in st.get("recent", []) if x != p and Path(x).exists()]
    st["recent"] = [p] + items[:max(0, limit - 1)]
    _write_state(st)


def session_paths(path: str | Path) -> dict[str, Path]:
    key = _file_key(path)
    folder = PDF_WORK_DIR / key
    folder.mkdir(parents=True, exist_ok=True)
    return {
        "folder": folder,
        "work": folder / "work.pdf",
        "redo": folder / "redo.pdf",
        "undo": folder / "undo.pdf",
        "meta": folder / "meta.json",
    }


def load_workspace(path: str | Path) -> tuple[Path, bool]:
    src = Path(path).resolve()
    sp = session_paths(src)
    restored = False
    if sp["work"].exists() and sp["meta"].exists():
        try:
            meta = json.loads(sp["meta"].read_text(encoding="utf-8"))
            restored = bool(meta.get("dirty")) and Path(meta.get("source", "")) == src
        except Exception:
            restored = False
    if not sp["work"].exists() or not restored:
        shutil.copy2(src, sp["work"])
        save_workspace_meta(src, dirty=False)
    remember_recent(src)
    return sp["work"], restored


def save_workspace_meta(source: str | Path, dirty: bool = False) -> None:
    src = Path(source).resolve()
    sp = session_paths(src)
    data = {
        "source": str(src),
        "dirty": bool(dirty),
        "updated_at": _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    sp["meta"].write_text(json.dumps(data, indent=2), encoding="utf-8")


def snapshot_undo(source: str | Path) -> None:
    src = Path(source).resolve()
    sp = session_paths(src)
    if sp["work"].exists():
        shutil.copy2(sp["work"], sp["undo"])


def can_undo(source: str | Path) -> bool:
    return session_paths(source)["undo"].exists()


def can_redo(source: str | Path) -> bool:
    return session_paths(source)["redo"].exists()


def undo(source: str | Path) -> bool:
    src = Path(source).resolve()
    sp = session_paths(src)
    if not sp["undo"].exists() or not sp["work"].exists():
        return False
    shutil.copy2(sp["work"], sp["redo"])
    shutil.copy2(sp["undo"], sp["work"])
    save_workspace_meta(src, dirty=True)
    return True


def redo(source: str | Path) -> bool:
    src = Path(source).resolve()
    sp = session_paths(src)
    if not sp["redo"].exists() or not sp["work"].exists():
        return False
    shutil.copy2(sp["work"], sp["undo"])
    shutil.copy2(sp["redo"], sp["work"])
    save_workspace_meta(src, dirty=True)
    return True


def clear_workspace(source: str | Path) -> None:
    src = Path(source).resolve()
    sp = session_paths(src)
    for p in (sp["undo"], sp["redo"]):
        try:
            p.unlink()
        except OSError:
            pass
    save_workspace_meta(src, dirty=False)


def commit_workspace(source: str | Path, dest: str | Path | None = None) -> Path:
    src = Path(source).resolve()
    sp = session_paths(src)
    out = Path(dest).resolve() if dest else src
    tmp = out.with_suffix(out.suffix + ".tmp")
    shutil.copy2(sp["work"], tmp)
    tmp.replace(out)
    if out == src:
        save_workspace_meta(src, dirty=False)
    remember_recent(out)
    return out


def page_count(path: str | Path) -> int:
    return len(PdfReader(str(path)).pages)


def page_sizes(path: str | Path) -> list[tuple[float, float]]:
    out = []
    for p in PdfReader(str(path)).pages:
        out.append((float(p.mediabox.width), float(p.mediabox.height)))
    return out


def render_page(path: str | Path, page_index: int, scale: float = 1.3,
                rotation: int = 0):
    pdf = pdfium.PdfDocument(str(path))
    page = pdf[page_index]
    bmp = page.render(scale=max(0.2, float(scale)), rotation=int(rotation) % 360)
    return bmp.to_pil()


def export_page_image(path: str | Path, page_index: int, out_path: str | Path,
                      scale: float = 2.0, rotation: int = 0) -> Path:
    img = render_page(path, page_index, scale=scale, rotation=rotation)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    if out.suffix.lower() in {".jpg", ".jpeg"}:
        img = img.convert("RGB")
    img.save(out)
    return out


def export_all_images(path: str | Path, out_dir: str | Path, fmt: str = "png",
                      scale: float = 2.0) -> list[Path]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    n = page_count(path)
    out = []
    ext = ".jpg" if fmt.lower() in {"jpg", "jpeg"} else ".png"
    for i in range(n):
        out.append(export_page_image(path, i, out_dir / f"page_{i + 1:03d}{ext}", scale=scale))
    return out


def extract_text(path: str | Path) -> list[str]:
    out = []
    r = PdfReader(str(path))
    for p in r.pages:
        try:
            out.append((p.extract_text() or "").strip())
        except Exception:
            out.append("")
    return out


def search_text(path: str | Path, query: str, limit: int = 200) -> list[dict[str, Any]]:
    q = str(query or "").strip().lower()
    if not q:
        return []
    out = []
    for i, txt in enumerate(extract_text(path), start=1):
        low = txt.lower()
        pos = low.find(q)
        if pos < 0:
            continue
        start = max(0, pos - 70)
        end = min(len(txt), pos + len(q) + 160)
        snippet = " ".join(txt[start:end].split())
        if start > 0:
            snippet = "… " + snippet
        if end < len(txt):
            snippet += " …"
        out.append({"page": i, "snippet": snippet or f"Match on page {i}"})
        if len(out) >= limit:
            break
    return out


def _reader(path: str | Path) -> PdfReader:
    return PdfReader(str(path))


def _write_pages(reader: PdfReader, page_indexes: Iterable[int], out_path: str | Path,
                 rotate_map: dict[int, int] | None = None) -> Path:
    writer = PdfWriter()
    rotate_map = rotate_map or {}
    for i in page_indexes:
        page = reader.pages[int(i)]
        angle = int(rotate_map.get(int(i), 0) or 0)
        if angle:
            page = page.rotate(angle)
        writer.add_page(page)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "wb") as fh:
        writer.write(fh)
    return out


def _page_indexes_from_text(text: str, total: int) -> list[int]:
    picked: list[int] = []
    for part in str(text or "").replace(";", ",").split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = [x.strip() for x in part.split("-", 1)]
            if a.isdigit() and b.isdigit():
                s, e = max(1, int(a)), min(total, int(b))
                if s <= e:
                    picked.extend(range(s - 1, e))
        elif part.isdigit():
            n = int(part)
            if 1 <= n <= total:
                picked.append(n - 1)
    out: list[int] = []
    for i in picked:
        if i not in out:
            out.append(i)
    return out


def parse_page_range(text: str, total: int) -> list[int]:
    idx = _page_indexes_from_text(text, total)
    return idx or list(range(total))


def merge_pdfs(paths: list[str | Path], out_path: str | Path) -> Path:
    writer = PdfWriter()
    for p in paths:
        for page in PdfReader(str(p)).pages:
            writer.add_page(page)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "wb") as fh:
        writer.write(fh)
    return out


def split_pdf(path: str | Path, page_ranges: list[str], out_dir: str | Path) -> list[Path]:
    reader = _reader(path)
    total = len(reader.pages)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out = []
    stem = _safe(Path(path).stem, "split")
    for rng in page_ranges:
        idx = parse_page_range(rng, total)
        label = _safe(rng.replace(",", "_"), "all")
        out.append(_write_pages(reader, idx, out_dir / f"{stem}_{label}.pdf"))
    return out


def reorder_pdf(path: str | Path, order: list[int], out_path: str | Path) -> Path:
    return _write_pages(_reader(path), order, out_path)


def duplicate_pages(path: str | Path, page_indexes: list[int], out_path: str | Path) -> Path:
    reader = _reader(path)
    order = list(range(len(reader.pages)))
    extra = sorted(int(i) for i in page_indexes)
    shift = 0
    for idx in extra:
        order.insert(idx + 1 + shift, idx)
        shift += 1
    return _write_pages(reader, order, out_path)


def delete_pages(path: str | Path, page_indexes: list[int], out_path: str | Path) -> Path:
    reader = _reader(path)
    skip = {int(i) for i in page_indexes}
    keep = [i for i in range(len(reader.pages)) if i not in skip]
    if not keep:
        raise ValueError("A PDF must keep at least one page.")
    return _write_pages(reader, keep, out_path)


def extract_pages(path: str | Path, page_indexes: list[int], out_path: str | Path) -> Path:
    return _write_pages(_reader(path), page_indexes, out_path)


def rotate_pages(path: str | Path, page_indexes: list[int], angle: int,
                 out_path: str | Path) -> Path:
    reader = _reader(path)
    ang = int(angle) % 360
    rotate_map = {int(i): ang for i in page_indexes}
    return _write_pages(reader, range(len(reader.pages)), out_path, rotate_map)


def duplicate_file(path: str | Path, dest_dir: str | Path | None = None) -> Path:
    src = Path(path)
    folder = Path(dest_dir or src.parent)
    folder.mkdir(parents=True, exist_ok=True)
    cand = folder / f"{src.stem} copy{src.suffix}"
    n = 2
    while cand.exists():
        cand = folder / f"{src.stem} copy {n}{src.suffix}"
        n += 1
    shutil.copy2(src, cand)
    return cand


def rename_file(path: str | Path, new_name: str) -> Path:
    src = Path(path)
    name = _safe(new_name, src.stem)
    if not name.lower().endswith(src.suffix.lower()):
        name += src.suffix
    dest = src.with_name(name)
    if dest == src:
        return dest
    if dest.exists():
        raise FileExistsError(dest)
    src.rename(dest)
    return dest


def _copy_name(path: str | Path, suffix: str) -> Path:
    p = Path(path)
    PDF_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return PDF_EXPORT_DIR / f"{_safe(p.stem)}_{_now_tag()}{suffix}"


def convert_to_docx(path: str | Path, out_path: str | Path | None = None) -> Path:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not available in this build.")
    out = Path(out_path or _copy_name(path, ".docx"))
    doc = DocxDocument()
    doc.add_heading(Path(path).name, 0)
    for i, txt in enumerate(extract_text(path), start=1):
        doc.add_heading(f"Page {i}", level=1)
        doc.add_paragraph(txt or "(no searchable text on this page)")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def convert_to_xlsx(path: str | Path, out_path: str | Path | None = None) -> Path:
    out = Path(out_path or _copy_name(path, ".xlsx"))
    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Text"
    ws.append(["Page", "Text"])
    for i, txt in enumerate(extract_text(path), start=1):
        ws.append([i, txt])
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 120
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


def convert_to_txt(path: str | Path, out_path: str | Path | None = None) -> Path:
    out = Path(out_path or _copy_name(path, ".txt"))
    parts = []
    for i, txt in enumerate(extract_text(path), start=1):
        parts.append(f"PAGE {i}\n{txt}\n")
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def convert_to_html(path: str | Path, out_path: str | Path | None = None) -> Path:
    out = Path(out_path or _copy_name(path, ".html"))
    body = [f"<h1>{html.escape(Path(path).name)}</h1>"]
    for i, txt in enumerate(extract_text(path), start=1):
        body.append(f"<h2>Page {i}</h2><pre>{html.escape(txt)}</pre>")
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("<html><body>" + "\n".join(body) + "</body></html>", encoding="utf-8")
    return out


def permissions_flag(allow_print: bool = True, allow_copy: bool = True,
                     allow_modify: bool = True) -> Any:
    flag = UAP(0)
    if allow_print:
        flag |= UAP.PRINT | UAP.PRINT_TO_REPRESENTATION
    if allow_copy:
        flag |= UAP.EXTRACT | UAP.EXTRACT_TEXT_AND_GRAPHICS
    if allow_modify:
        flag |= UAP.MODIFY | UAP.ADD_OR_MODIFY | UAP.FILL_FORM_FIELDS | UAP.ASSEMBLE_DOC
    return flag


def protect_pdf(path: str | Path, out_path: str | Path, user_password: str,
                owner_password: str = "", allow_print: bool = True,
                allow_copy: bool = True, allow_modify: bool = False) -> Path:
    reader = _reader(path)
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    perms = permissions_flag(allow_print, allow_copy, allow_modify)
    try:
        writer.encrypt(user_password=user_password,
                       owner_password=owner_password or user_password,
                       permissions_flag=perms,
                       algorithm="AES-256")
    except Exception:
        writer.encrypt(user_password=user_password,
                       owner_password=owner_password or user_password,
                       permissions_flag=perms,
                       use_128bit=True)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "wb") as fh:
        writer.write(fh)
    return out


def _pct(value: float, total: float) -> float:
    return float(value or 0) / 100.0 * float(total)


def _overlay_page(width: float, height: float, anns: list[dict[str, Any]]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(width, height))
    for a in anns:
        kind = str(a.get("type") or "text").lower()
        x = _pct(float(a.get("x") or 5), width)
        w = max(8.0, _pct(float(a.get("w") or 18), width))
        h = max(8.0, _pct(float(a.get("h") or 5), height))
        y_top = _pct(float(a.get("y") or 5), height)
        y = height - y_top - h
        fill = str(a.get("color") or "#ffbf00")
        text = str(a.get("text") or "")
        try:
            stroke = colors.HexColor(fill) if fill.startswith("#") else colors.HexColor("#ffbf00")
        except Exception:
            stroke = colors.HexColor("#ffbf00")
        if kind == "highlight":
            c.setFillColor(colors.Color(stroke.red, stroke.green, stroke.blue, alpha=0.28))
            c.rect(x, y, w, h, fill=1, stroke=0)
        elif kind == "underline":
            c.setStrokeColor(stroke)
            c.setLineWidth(max(1, h / 8))
            c.line(x, y + 2, x + w, y + 2)
        elif kind == "line":
            c.setStrokeColor(stroke)
            c.setLineWidth(max(1, min(w, h) / 20))
            c.line(x, y, x + w, y + h)
        elif kind == "box":
            c.setStrokeColor(stroke)
            c.setLineWidth(1.4)
            c.rect(x, y, w, h, fill=0, stroke=1)
        elif kind == "stamp":
            c.setFillColor(stroke)
            c.roundRect(x, y, w, h, 6, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", max(10, min(22, h * 0.45)))
            c.drawCentredString(x + w / 2, y + h * 0.35, text or "APPROVED")
        elif kind == "note":
            c.setFillColor(colors.Color(1, 0.97, 0.72, alpha=0.96))
            c.roundRect(x, y, w, h, 6, fill=1, stroke=1)
            para = Paragraph(html.escape(text or "Note"), _STYLES["BodyText"])
            para.wrapOn(c, w - 8, h - 8)
            para.drawOn(c, x + 4, y + 4)
        elif kind == "signature":
            img = str(a.get("image_path") or "")
            if img and Path(img).exists():
                c.drawImage(img, x, y, w, h, preserveAspectRatio=True, mask="auto")
            if text:
                c.setFont("Helvetica", max(8, min(12, h * 0.18)))
                c.setFillColor(colors.black)
                c.drawString(x, max(4, y - 12), text)
        else:  # text
            para = Paragraph(html.escape(text or "Text"), _STYLES["BodyText"])
            para.wrapOn(c, w, h + 32)
            para.drawOn(c, x, y)
    c.showPage()
    c.save()
    return buf.getvalue()


def annotate_pdf(path: str | Path, annotations: list[dict[str, Any]],
                 out_path: str | Path) -> Path:
    reader = _reader(path)
    writer = PdfWriter()
    by_page: dict[int, list[dict[str, Any]]] = {}
    for ann in annotations:
        page = max(1, int(ann.get("page") or 1)) - 1
        by_page.setdefault(page, []).append(ann)
    for i, page in enumerate(reader.pages):
        if i in by_page:
            overlay_bytes = _overlay_page(float(page.mediabox.width), float(page.mediabox.height),
                                          by_page[i])
            overlay_reader = PdfReader(io.BytesIO(overlay_bytes))
            page = page
            page.merge_page(overlay_reader.pages[0])
        writer.add_page(page)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "wb") as fh:
        writer.write(fh)
    return out


def paste_signature(path: str | Path, image_path: str | Path, page: int,
                    x: float = 70, y: float = 82, w: float = 22, h: float = 10,
                    caption: str = "") -> Path:
    tmp = _copy_name(path, ".signed.pdf")
    return annotate_pdf(path, [{
        "type": "signature", "page": page, "x": x, "y": y, "w": w, "h": h,
        "image_path": str(image_path), "text": caption,
    }], tmp)


def page_summary(path: str | Path) -> list[dict[str, Any]]:
    out = []
    for i, (w, h) in enumerate(page_sizes(path), start=1):
        out.append({"page": i, "width": round(w, 1), "height": round(h, 1)})
    return out


def default_export_path(path: str | Path, suffix: str) -> Path:
    return _copy_name(path, suffix)


def merge_to_workspace(source: str | Path, extra_paths: list[str | Path]) -> Path:
    src = Path(source).resolve()
    sp = session_paths(src)
    snapshot_undo(src)
    merge_pdfs([sp["work"], *extra_paths], sp["work"])
    save_workspace_meta(src, dirty=True)
    return sp["work"]


def replace_workspace(source: str | Path, maker) -> Path:
    src = Path(source).resolve()
    sp = session_paths(src)
    snapshot_undo(src)
    tmp = sp["folder"] / f"work_{_now_tag()}.pdf"
    maker(sp["work"], tmp)
    tmp.replace(sp["work"])
    save_workspace_meta(src, dirty=True)
    return sp["work"]
